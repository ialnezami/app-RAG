"""
Advanced document processing with OCR, language detection, and quality assessment.
"""
import os
import logging
from typing import Dict, List, Optional, Tuple, Any
from pathlib import Path
import mimetypes
import tempfile

# Optional imports for advanced features
try:
    import cv2
    import numpy as np
    HAS_CV2 = True
except ImportError:
    HAS_CV2 = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    import pytesseract
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

try:
    import langdetect
    HAS_LANGDETECT = True
except ImportError:
    HAS_LANGDETECT = False

logger = logging.getLogger(__name__)

class AdvancedDocumentProcessor:
    """Advanced document processing with OCR and analysis capabilities."""
    
    def __init__(self):
        self.supported_image_types = {'.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif'}
        self.supported_text_types = {'.txt', '.md', '.csv'}
        self.supported_pdf_types = {'.pdf'}
        self.supported_doc_types = {'.docx', '.doc'}
        
    async def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        Advanced document processing with OCR, language detection, and quality assessment.
        
        Args:
            file_path: Path to the uploaded file
            filename: Original filename
            
        Returns:
            Dictionary with processing results
        """
        try:
            file_ext = Path(filename).suffix.lower()
            mime_type, _ = mimetypes.guess_type(filename)
            
            result = {
                'filename': filename,
                'file_path': file_path,
                'mime_type': mime_type,
                'file_extension': file_ext,
                'processing_status': 'success',
                'extracted_text': '',
                'language': 'unknown',
                'quality_score': 0.0,
                'metadata': {},
                'ocr_confidence': 0.0,
                'has_images': False,
                'page_count': 0,
                'word_count': 0,
                'character_count': 0,
                'processing_method': 'unknown'
            }
            
            # Process based on file type
            if file_ext in self.supported_image_types:
                result.update(await self._process_image(file_path))
                result['processing_method'] = 'ocr'
            elif file_ext in self.supported_pdf_types:
                result.update(await self._process_pdf(file_path))
                result['processing_method'] = 'pdf_extraction'
            elif file_ext in self.supported_doc_types:
                result.update(await self._process_document_file(file_path))
                result['processing_method'] = 'document_extraction'
            elif file_ext in self.supported_text_types:
                result.update(await self._process_text_file(file_path))
                result['processing_method'] = 'text_extraction'
            else:
                result['processing_status'] = 'unsupported'
                result['error'] = f'Unsupported file type: {file_ext}'
                return result
            
            # Post-processing analysis
            if result['extracted_text']:
                result.update(await self._analyze_text(result['extracted_text']))
            
            # Calculate overall quality score
            result['quality_score'] = self._calculate_quality_score(result)
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            return {
                'filename': filename,
                'processing_status': 'error',
                'error': str(e),
                'quality_score': 0.0
            }
    
    async def _process_image(self, file_path: str) -> Dict[str, Any]:
        """Process image files with OCR."""
        result = {
            'extracted_text': '',
            'ocr_confidence': 0.0,
            'has_images': True,
            'page_count': 1
        }
        
        if not HAS_TESSERACT or not HAS_PIL:
            result['error'] = 'OCR dependencies not available (tesseract, PIL)'
            return result
        
        try:
            # Open and preprocess image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Enhance image quality for better OCR
            if HAS_CV2:
                # Convert PIL to OpenCV
                cv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
                
                # Apply image preprocessing
                cv_image = self._enhance_image_for_ocr(cv_image)
                
                # Convert back to PIL
                image = Image.fromarray(cv2.cvtColor(cv_image, cv2.COLOR_BGR2RGB))
            
            # Perform OCR
            ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            
            # Extract text and calculate confidence
            words = []
            confidences = []
            
            for i, word in enumerate(ocr_data['text']):
                if word.strip():
                    words.append(word)
                    confidences.append(ocr_data['conf'][i])
            
            result['extracted_text'] = ' '.join(words)
            result['ocr_confidence'] = np.mean(confidences) if confidences else 0.0
            
            # Extract additional metadata
            result['metadata'] = {
                'image_size': image.size,
                'image_mode': image.mode,
                'words_detected': len(words),
                'avg_confidence': result['ocr_confidence']
            }
            
        except Exception as e:
            result['error'] = f'OCR processing failed: {str(e)}'
            logger.error(f"OCR processing error: {e}")
        
        return result
    
    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF files with text extraction and OCR fallback."""
        result = {
            'extracted_text': '',
            'page_count': 0,
            'has_images': False
        }
        
        try:
            # Try text extraction first (faster)
            from PyPDF2 import PdfReader
            
            reader = PdfReader(file_path)
            result['page_count'] = len(reader.pages)
            
            text_content = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(page_text)
            
            result['extracted_text'] = '\n'.join(text_content)
            
            # If no text extracted, try OCR (if available)
            if not result['extracted_text'].strip() and HAS_TESSERACT:
                result.update(await self._ocr_pdf(file_path))
            
            # Check for images
            result['has_images'] = self._pdf_has_images(reader)
            
        except Exception as e:
            result['error'] = f'PDF processing failed: {str(e)}'
            logger.error(f"PDF processing error: {e}")
        
        return result
    
    async def _process_document_file(self, file_path: str) -> Dict[str, Any]:
        """Process Word documents."""
        result = {
            'extracted_text': '',
            'page_count': 1,
            'has_images': False
        }
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            result['extracted_text'] = '\n'.join(paragraphs)
            
            # Check for images/shapes
            result['has_images'] = len(doc.inline_shapes) > 0
            
            # Metadata
            result['metadata'] = {
                'paragraphs': len(doc.paragraphs),
                'inline_shapes': len(doc.inline_shapes),
                'tables': len(doc.tables)
            }
            
        except Exception as e:
            result['error'] = f'Document processing failed: {str(e)}'
            logger.error(f"Document processing error: {e}")
        
        return result
    
    async def _process_text_file(self, file_path: str) -> Dict[str, Any]:
        """Process plain text files."""
        result = {
            'extracted_text': '',
            'page_count': 1,
            'has_images': False
        }
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        result['extracted_text'] = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not result['extracted_text']:
                result['error'] = 'Could not decode text file with any supported encoding'
            
        except Exception as e:
            result['error'] = f'Text processing failed: {str(e)}'
            logger.error(f"Text processing error: {e}")
        
        return result
    
    async def _analyze_text(self, text: str) -> Dict[str, Any]:
        """Analyze extracted text for language, quality, etc."""
        analysis = {
            'language': 'unknown',
            'word_count': 0,
            'character_count': len(text),
            'sentences': 0,
            'readability_score': 0.0
        }
        
        try:
            # Basic text statistics
            words = text.split()
            analysis['word_count'] = len(words)
            
            sentences = text.split('.')
            analysis['sentences'] = len([s for s in sentences if s.strip()])
            
            # Language detection
            if HAS_LANGDETECT and text.strip():
                try:
                    analysis['language'] = langdetect.detect(text)
                except:
                    analysis['language'] = 'unknown'
            
            # Simple readability score (Flesch-like)
            if analysis['sentences'] > 0 and analysis['word_count'] > 0:
                avg_words_per_sentence = analysis['word_count'] / analysis['sentences']
                avg_chars_per_word = analysis['character_count'] / analysis['word_count']
                
                # Simple readability metric (lower is better)
                analysis['readability_score'] = min(100, max(0, 
                    100 - (avg_words_per_sentence * 2) - (avg_chars_per_word * 5)
                ))
            
        except Exception as e:
            logger.error(f"Text analysis error: {e}")
        
        return analysis
    
    def _enhance_image_for_ocr(self, image: np.ndarray) -> np.ndarray:
        """Enhance image quality for better OCR results."""
        if not HAS_CV2:
            return image
        
        try:
            # Convert to grayscale
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Apply Gaussian blur to reduce noise
            blurred = cv2.GaussianBlur(gray, (5, 5), 0)
            
            # Apply threshold to get binary image
            _, thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # Apply morphological operations to clean up
            kernel = np.ones((2, 2), np.uint8)
            cleaned = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
            
            # Convert back to BGR for consistency
            return cv2.cvtColor(cleaned, cv2.COLOR_GRAY2BGR)
            
        except Exception as e:
            logger.error(f"Image enhancement error: {e}")
            return image
    
    async def _ocr_pdf(self, file_path: str) -> Dict[str, Any]:
        """Perform OCR on PDF pages."""
        result = {
            'extracted_text': '',
            'ocr_confidence': 0.0
        }
        
        # This would require pdf2image library
        # For now, return empty result
        result['error'] = 'PDF OCR not implemented (requires pdf2image)'
        return result
    
    def _pdf_has_images(self, reader) -> bool:
        """Check if PDF contains images."""
        try:
            for page in reader.pages:
                if '/XObject' in page['/Resources']:
                    xObject = page['/Resources']['/XObject'].get_object()
                    for obj in xObject:
                        if xObject[obj]['/Subtype'] == '/Image':
                            return True
            return False
        except:
            return False
    
    def _calculate_quality_score(self, result: Dict[str, Any]) -> float:
        """Calculate overall document quality score (0-100)."""
        score = 0.0
        
        try:
            # Base score for successful processing
            if result.get('processing_status') == 'success':
                score += 20
            
            # Text content quality
            text_length = len(result.get('extracted_text', ''))
            if text_length > 0:
                score += min(30, text_length / 100)  # Up to 30 points for text length
            
            # Word count bonus
            word_count = result.get('word_count', 0)
            if word_count > 0:
                score += min(20, word_count / 50)  # Up to 20 points for word count
            
            # OCR confidence (if applicable)
            ocr_confidence = result.get('ocr_confidence', 0)
            if ocr_confidence > 0:
                score += (ocr_confidence / 100) * 15  # Up to 15 points for OCR quality
            
            # Language detection bonus
            if result.get('language') != 'unknown':
                score += 10
            
            # Readability bonus
            readability = result.get('readability_score', 0)
            if readability > 0:
                score += (readability / 100) * 5  # Up to 5 points for readability
            
            return min(100.0, max(0.0, score))
            
        except Exception as e:
            logger.error(f"Quality score calculation error: {e}")
            return 0.0
    
    def get_processing_capabilities(self) -> Dict[str, bool]:
        """Get available processing capabilities."""
        return {
            'ocr_available': HAS_TESSERACT and HAS_PIL,
            'image_processing': HAS_CV2 and HAS_PIL,
            'language_detection': HAS_LANGDETECT,
            'pdf_processing': True,  # PyPDF2 is in requirements
            'document_processing': True,  # python-docx is in requirements
            'supported_formats': list(
                self.supported_image_types | 
                self.supported_text_types | 
                self.supported_pdf_types | 
                self.supported_doc_types
            )
        }
