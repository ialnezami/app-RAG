"""
Text chunking and preprocessing utilities.
"""
import re
import uuid
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
import PyPDF2
import docx
import markdown
from bs4 import BeautifulSoup
import logging

logger = logging.getLogger(__name__)


@dataclass
class TextChunk:
    """Represents a chunk of text with metadata."""
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass
class DocumentMetadata:
    """Document metadata extracted during processing."""
    title: Optional[str] = None
    author: Optional[str] = None
    pages: Optional[int] = None
    word_count: Optional[int] = None
    file_size: Optional[int] = None
    mime_type: Optional[str] = None
    language: Optional[str] = None
    created_date: Optional[str] = None
    modified_date: Optional[str] = None
    custom_metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.custom_metadata is None:
            self.custom_metadata = {}


class TextChunker:
    """Handles text chunking with various strategies."""
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        min_chunk_size: int = 100,
        max_chunk_size: int = 2000,
        preserve_paragraphs: bool = True,
        preserve_sentences: bool = True
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        self.preserve_paragraphs = preserve_paragraphs
        self.preserve_sentences = preserve_sentences
    
    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[TextChunk]:
        """
        Chunk text into smaller pieces.
        
        Args:
            text: Input text to chunk
            metadata: Optional metadata to include with chunks
            
        Returns:
            List of TextChunk objects
        """
        if not text or not text.strip():
            return []
        
        # Clean and normalize text
        cleaned_text = self._clean_text(text)
        
        # Try different chunking strategies
        chunks = []
        
        if self.preserve_paragraphs:
            chunks = self._chunk_by_paragraphs(cleaned_text, metadata)
        
        if not chunks and self.preserve_sentences:
            chunks = self._chunk_by_sentences(cleaned_text, metadata)
        
        if not chunks:
            chunks = self._chunk_by_characters(cleaned_text, metadata)
        
        # Filter chunks by size
        filtered_chunks = self._filter_chunks_by_size(chunks)
        
        # Add metadata to chunks
        for i, chunk in enumerate(filtered_chunks):
            chunk.chunk_index = i
            if metadata:
                chunk.metadata.update(metadata)
        
        logger.info(f"Created {len(filtered_chunks)} chunks from text")
        return filtered_chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\r\n|\r', '\n', text)
        
        return text.strip()
    
    def _chunk_by_paragraphs(self, text: str, metadata: Dict[str, Any] = None) -> List[TextChunk]:
        """Chunk text by paragraphs."""
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        current_start = 0
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # If adding this paragraph would exceed chunk size, create a new chunk
            if current_chunk and len(current_chunk) + len(paragraph) + 2 > self.chunk_size:
                if current_chunk:
                    chunks.append(TextChunk(
                        content=current_chunk.strip(),
                        chunk_index=0,  # Will be set later
                        start_char=current_start,
                        end_char=current_start + len(current_chunk),
                        metadata=metadata or {}
                    ))
                    current_start += len(current_chunk) + 2
                    current_chunk = paragraph
                else:
                    current_chunk = paragraph
            else:
                if current_chunk:
                    current_chunk += '\n\n' + paragraph
                else:
                    current_chunk = paragraph
        
        # Add the last chunk
        if current_chunk:
            chunks.append(TextChunk(
                content=current_chunk.strip(),
                chunk_index=0,  # Will be set later
                start_char=current_start,
                end_char=current_start + len(current_chunk),
                metadata=metadata or {}
            ))
        
        return chunks
    
    def _chunk_by_sentences(self, text: str, metadata: Dict[str, Any] = None) -> List[TextChunk]:
        """Chunk text by sentences."""
        # Simple sentence splitting (can be improved with NLP libraries)
        sentences = re.split(r'[.!?]+\s+', text)
        chunks = []
        current_chunk = ""
        current_start = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # If adding this sentence would exceed chunk size, create a new chunk
            if current_chunk and len(current_chunk) + len(sentence) + 2 > self.chunk_size:
                if current_chunk:
                    chunks.append(TextChunk(
                        content=current_chunk.strip(),
                        chunk_index=0,  # Will be set later
                        start_char=current_start,
                        end_char=current_start + len(current_chunk),
                        metadata=metadata or {}
                    ))
                    current_start += len(current_chunk) + 2
                    current_chunk = sentence
                else:
                    current_chunk = sentence
            else:
                if current_chunk:
                    current_chunk += '. ' + sentence
                else:
                    current_chunk = sentence
        
        # Add the last chunk
        if current_chunk:
            chunks.append(TextChunk(
                content=current_chunk.strip(),
                chunk_index=0,  # Will be set later
                start_char=current_start,
                end_char=current_start + len(current_chunk),
                metadata=metadata or {}
            ))
        
        return chunks
    
    def _chunk_by_characters(self, text: str, metadata: Dict[str, Any] = None) -> List[TextChunk]:
        """Chunk text by character count with overlap."""
        chunks = []
        start = 0
        
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            
            # Try to break at word boundary
            if end < len(text):
                # Look for the last space within the chunk
                last_space = text.rfind(' ', start, end)
                if last_space > start + self.min_chunk_size:
                    end = last_space
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append(TextChunk(
                    content=chunk_text,
                    chunk_index=0,  # Will be set later
                    start_char=start,
                    end_char=end,
                    metadata=metadata or {}
                ))
            
            # Move start position with overlap
            start = max(start + self.chunk_size - self.chunk_overlap, end)
        
        return chunks
    
    def _filter_chunks_by_size(self, chunks: List[TextChunk]) -> List[TextChunk]:
        """Filter chunks by size constraints."""
        filtered = []
        
        for chunk in chunks:
            if self.min_chunk_size <= len(chunk.content) <= self.max_chunk_size:
                filtered.append(chunk)
            elif len(chunk.content) > self.max_chunk_size:
                # Split oversized chunks
                sub_chunks = self._split_oversized_chunk(chunk)
                filtered.extend(sub_chunks)
        
        return filtered
    
    def _split_oversized_chunk(self, chunk: TextChunk) -> List[TextChunk]:
        """Split an oversized chunk into smaller pieces."""
        content = chunk.content
        chunks = []
        start = 0
        
        while start < len(content):
            end = min(start + self.max_chunk_size, len(content))
            
            # Try to break at word boundary
            if end < len(content):
                last_space = content.rfind(' ', start, end)
                if last_space > start + self.min_chunk_size:
                    end = last_space
            
            chunk_text = content[start:end].strip()
            if chunk_text:
                chunks.append(TextChunk(
                    content=chunk_text,
                    chunk_index=0,  # Will be set later
                    start_char=chunk.start_char + start,
                    end_char=chunk.start_char + end,
                    metadata=chunk.metadata.copy()
                ))
            
            start = end
        
        return chunks


class DocumentProcessor:
    """Handles document processing and text extraction."""
    
    def __init__(self):
        self.chunker = TextChunker()
    
    def process_document(
        self,
        file_path: str,
        mime_type: str,
        metadata: Dict[str, Any] = None
    ) -> Tuple[str, DocumentMetadata, List[TextChunk]]:
        """
        Process a document and extract text and chunks.
        
        Args:
            file_path: Path to the document file
            mime_type: MIME type of the document
            metadata: Additional metadata
            
        Returns:
            Tuple of (extracted_text, document_metadata, text_chunks)
        """
        logger.info(f"Processing document: {file_path} ({mime_type})")
        
        # Extract text based on file type
        if mime_type == "application/pdf":
            text, doc_metadata = self._extract_pdf_text(file_path)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text, doc_metadata = self._extract_docx_text(file_path)
        elif mime_type == "text/markdown":
            text, doc_metadata = self._extract_markdown_text(file_path)
        elif mime_type == "text/plain":
            text, doc_metadata = self._extract_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")
        
        # Merge with provided metadata
        if metadata:
            doc_metadata.custom_metadata.update(metadata)
        
        # Chunk the extracted text
        chunks = self.chunker.chunk_text(text, {
            "file_path": file_path,
            "mime_type": mime_type,
            **doc_metadata.custom_metadata
        })
        
        logger.info(f"Extracted {len(text)} characters, created {len(chunks)} chunks")
        return text, doc_metadata, chunks
    
    def _extract_pdf_text(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text from PDF file."""
        text = ""
        metadata = DocumentMetadata()
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata.title = pdf_reader.metadata.get('/Title')
                    metadata.author = pdf_reader.metadata.get('/Author')
                    metadata.created_date = str(pdf_reader.metadata.get('/CreationDate', ''))
                    metadata.modified_date = str(pdf_reader.metadata.get('/ModDate', ''))
                
                metadata.pages = len(pdf_reader.pages)
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n'
                        metadata.custom_metadata[f'page_{page_num + 1}_text_length'] = len(page_text)
                
                metadata.word_count = len(text.split())
                metadata.mime_type = "application/pdf"
                
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise
        
        return text.strip(), metadata
    
    def _extract_docx_text(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text from DOCX file."""
        text = ""
        metadata = DocumentMetadata()
        
        try:
            doc = docx.Document(file_path)
            
            # Extract metadata
            if doc.core_properties:
                metadata.title = doc.core_properties.title
                metadata.author = doc.core_properties.author
                metadata.created_date = str(doc.core_properties.created) if doc.core_properties.created else None
                metadata.modified_date = str(doc.core_properties.modified) if doc.core_properties.modified else None
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + '\n'
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + '\n'
            
            metadata.word_count = len(text.split())
            metadata.mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
        
        return text.strip(), metadata
    
    def _extract_markdown_text(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text from Markdown file."""
        text = ""
        metadata = DocumentMetadata()
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Convert markdown to HTML then extract text
            html = markdown.markdown(content)
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()
            
            # Extract title from first heading
            first_heading = soup.find(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
            if first_heading:
                metadata.title = first_heading.get_text().strip()
            
            metadata.word_count = len(text.split())
            metadata.mime_type = "text/markdown"
            
        except Exception as e:
            logger.error(f"Error extracting Markdown text: {e}")
            raise
        
        return text.strip(), metadata
    
    def _extract_text_file(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text from plain text file."""
        text = ""
        metadata = DocumentMetadata()
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata.word_count = len(text.split())
            metadata.mime_type = "text/plain"
            
        except Exception as e:
            logger.error(f"Error extracting text file: {e}")
            raise
        
        return text.strip(), metadata
    
    def get_supported_mime_types(self) -> List[str]:
        """Get list of supported MIME types."""
        return [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/markdown",
            "text/plain"
        ]
    
    def is_supported_file_type(self, mime_type: str) -> bool:
        """Check if file type is supported."""
        return mime_type in self.get_supported_mime_types()


# Global instances
_chunker: Optional[TextChunker] = None
_processor: Optional[DocumentProcessor] = None


def get_chunker() -> TextChunker:
    """Get global text chunker instance."""
    global _chunker
    if _chunker is None:
        _chunker = TextChunker()
    return _chunker


def get_document_processor() -> DocumentProcessor:
    """Get global document processor instance."""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor
